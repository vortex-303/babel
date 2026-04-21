from __future__ import annotations

import io
import zipfile

from docx import Document as DocxDocument

from app.models import Chunk, Document, Job, JobStatus
from app.services.assemble import (
    assemble_docx,
    assemble_epub,
    assemble_markdown,
)


def _fixture() -> tuple[Job, Document, list[Chunk]]:
    doc = Document(
        id=1,
        filename="alice.pdf",
        mime_type="application/pdf",
        size_bytes=1000,
        page_count=5,
        word_count=200,
        token_count=300,
        stored_path="/tmp/alice.pdf",
    )
    job = Job(
        id=42,
        document_id=1,
        status=JobStatus.DONE,
        source_lang="en",
        target_lang="es-AR",
        model_adapter="llamacpp",
        model_name="translategemma-4b-it.Q4_K_M.gguf",
        chunk_count=2,
        translated_chunks=2,
    )
    chunks = [
        Chunk(
            id=1,
            job_id=42,
            idx=0,
            source_text="Alice was beginning.",
            translated_text="Alicia empezaba.\n\nEstaba aburrida.",
            token_count=10,
        ),
        Chunk(
            id=2,
            job_id=42,
            idx=1,
            source_text="Suddenly a rabbit.",
            translated_text="De repente, un conejo.",
            token_count=8,
        ),
    ]
    return job, doc, chunks


def test_markdown_assembly():
    job, doc, chunks = _fixture()
    out = assemble_markdown(job, doc, chunks)
    text = out.content.decode()

    assert "alice.pdf" in text
    assert "en → es-AR" in text
    assert "Alicia empezaba." in text
    assert "De repente, un conejo." in text
    assert out.mime_type.startswith("text/markdown")
    assert out.filename.endswith(".md")
    assert "es-AR" in out.filename


def test_docx_assembly_produces_valid_document(tmp_path):
    job, doc, chunks = _fixture()
    out = assemble_docx(job, doc, chunks)

    assert out.mime_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert out.filename.endswith(".docx")

    path = tmp_path / "out.docx"
    path.write_bytes(out.content)
    reopened = DocxDocument(str(path))
    joined = "\n".join(p.text for p in reopened.paragraphs)
    assert "Alicia empezaba." in joined
    assert "De repente, un conejo." in joined


def test_epub_assembly_is_valid_zip():
    job, doc, chunks = _fixture()
    out = assemble_epub(job, doc, chunks)

    assert out.mime_type == "application/epub+zip"
    assert out.filename.endswith(".epub")

    with zipfile.ZipFile(io.BytesIO(out.content)) as zf:
        names = zf.namelist()
        assert any(n.endswith("chapter.xhtml") for n in names)
        assert "mimetype" in names or any(n == "mimetype" for n in names)
        chapter = [n for n in names if n.endswith("chapter.xhtml")][0]
        html = zf.read(chapter).decode()
        assert "Alicia empezaba." in html
        assert "De repente, un conejo." in html


def test_markdown_skips_empty_chunks():
    job, doc, chunks = _fixture()
    chunks.append(
        Chunk(
            id=3,
            job_id=42,
            idx=2,
            source_text="x",
            translated_text=None,
            token_count=0,
        )
    )
    chunks.append(
        Chunk(
            id=4,
            job_id=42,
            idx=3,
            source_text="y",
            translated_text="   ",
            token_count=0,
        )
    )
    out = assemble_markdown(job, doc, chunks)
    text = out.content.decode()
    assert text.count("Alicia empezaba.") == 1
    assert text.rstrip().endswith("De repente, un conejo.")
