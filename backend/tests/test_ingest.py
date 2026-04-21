from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from app.services.ingest import ingest, supported_extensions


def test_supported_extensions_includes_docx():
    assert ".docx" in supported_extensions()
    assert ".pdf" in supported_extensions()
    assert ".epub" in supported_extensions()


def test_ingest_txt(tmp_path: Path):
    p = tmp_path / "note.txt"
    p.write_text("First paragraph.\n\nSecond paragraph here.\n\nThird one.\n")
    result = ingest(p)
    assert len(result.chapters) == 1
    assert result.chapters[0].paragraphs == [
        "First paragraph.",
        "Second paragraph here.",
        "Third one.",
    ]
    assert result.word_count > 0


def test_ingest_md_treated_as_text(tmp_path: Path):
    p = tmp_path / "readme.md"
    p.write_text("# Title\n\nBody line one.\n\nBody line two.")
    result = ingest(p)
    assert result.page_count == 1
    assert len(result.chapters[0].paragraphs) == 3


def test_ingest_docx_heading_splits_chapters(tmp_path: Path):
    docx = DocxDocument()
    docx.add_heading("Chapter One", level=1)
    docx.add_paragraph("Opening paragraph of chapter one.")
    docx.add_paragraph("Second paragraph of chapter one.")
    docx.add_heading("Chapter Two", level=1)
    docx.add_paragraph("Only paragraph of chapter two.")
    p = tmp_path / "book.docx"
    docx.save(str(p))

    result = ingest(p)
    assert len(result.chapters) == 2
    assert result.chapters[0].title == "Chapter One"
    assert result.chapters[1].title == "Chapter Two"
    assert len(result.chapters[0].paragraphs) == 2
    assert len(result.chapters[1].paragraphs) == 1


def test_ingest_docx_no_headings(tmp_path: Path):
    docx = DocxDocument()
    docx.add_paragraph("Just a plain paragraph.")
    docx.add_paragraph("Another one.")
    p = tmp_path / "plain.docx"
    docx.save(str(p))

    result = ingest(p)
    assert len(result.chapters) == 1
    assert result.chapters[0].title == "plain"
    assert len(result.chapters[0].paragraphs) == 2


def test_ingest_rejects_unknown_ext(tmp_path: Path):
    p = tmp_path / "x.bin"
    p.write_bytes(b"\x00\x01")
    import pytest

    with pytest.raises(ValueError):
        ingest(p)
