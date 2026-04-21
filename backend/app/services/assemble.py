from __future__ import annotations

import io
import re
from collections.abc import Sequence
from dataclasses import dataclass

import ebooklib
from docx import Document as DocxDocument
from ebooklib import epub

from app.models import Chunk, Document, Job


@dataclass
class AssembledOutput:
    content: bytes
    mime_type: str
    filename: str


def _stem(doc: Document) -> str:
    name = doc.filename.rsplit(".", 1)[0] if "." in doc.filename else doc.filename
    return re.sub(r"[^A-Za-z0-9_-]+", "_", name).strip("_") or "translation"


def _joined_text(chunks: Sequence[Chunk]) -> str:
    """Concatenate translated chunk text. Assumes non-overlapping chunks."""
    parts: list[str] = []
    for c in chunks:
        text = (c.translated_text or "").strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _paragraphs(text: str) -> list[str]:
    return [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]


def assemble_markdown(job: Job, doc: Document, chunks: Sequence[Chunk]) -> AssembledOutput:
    body = _joined_text(chunks)
    header = (
        f"# {doc.filename}\n\n"
        f"*{job.source_lang} → {job.target_lang} · "
        f"translated with {job.model_adapter} ({job.model_name}) · "
        f"babel*\n\n"
    )
    content = (header + body).encode("utf-8")
    return AssembledOutput(
        content=content,
        mime_type="text/markdown; charset=utf-8",
        filename=f"{_stem(doc)}.{job.target_lang}.md",
    )


def assemble_docx(job: Job, doc: Document, chunks: Sequence[Chunk]) -> AssembledOutput:
    out = DocxDocument()
    title = out.add_paragraph()
    run = title.add_run(doc.filename)
    run.bold = True
    run.font.size = None  # default
    out.add_paragraph(
        f"{job.source_lang} → {job.target_lang} · translated with "
        f"{job.model_adapter} ({job.model_name}) · babel"
    ).italic = True
    out.add_paragraph("")

    for para in _paragraphs(_joined_text(chunks)):
        out.add_paragraph(para)

    buf = io.BytesIO()
    out.save(buf)
    return AssembledOutput(
        content=buf.getvalue(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{_stem(doc)}.{job.target_lang}.docx",
    )


def assemble_epub(job: Job, doc: Document, chunks: Sequence[Chunk]) -> AssembledOutput:
    book = epub.EpubBook()
    book.set_identifier(f"babel-job-{job.id}")
    book.set_title(f"{doc.filename} — {job.target_lang}")
    book.set_language(job.target_lang.split("-")[0])
    book.add_author("babel translation pipeline")

    chapter = epub.EpubHtml(
        title=doc.filename,
        file_name="chapter.xhtml",
        lang=job.target_lang.split("-")[0],
    )
    paragraphs_html = "\n".join(
        f"<p>{_escape_html(p)}</p>" for p in _paragraphs(_joined_text(chunks))
    )
    chapter.content = (
        "<html><head><title>"
        f"{_escape_html(doc.filename)}"
        "</title></head><body>"
        f"<h1>{_escape_html(doc.filename)}</h1>"
        f"<p><em>{job.source_lang} → {job.target_lang} · babel</em></p>"
        f"{paragraphs_html}"
        "</body></html>"
    )
    book.add_item(chapter)
    book.toc = (epub.Link("chapter.xhtml", doc.filename, "chapter"),)
    book.spine = ["nav", chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    buf = io.BytesIO()
    epub.write_epub(buf, book, {})
    return AssembledOutput(
        content=buf.getvalue(),
        mime_type="application/epub+zip",
        filename=f"{_stem(doc)}.{job.target_lang}.epub",
    )


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


ASSEMBLERS = {
    "md": assemble_markdown,
    "docx": assemble_docx,
    "epub": assemble_epub,
}


__all__ = ["ASSEMBLERS", "AssembledOutput", "assemble_docx", "assemble_epub", "assemble_markdown"]

# Keep ebooklib import referenced so static analyzers don't drop it.
_ = ebooklib
