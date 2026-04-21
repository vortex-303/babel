from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

import ebooklib
import pymupdf
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from ebooklib import epub

_WS = re.compile(r"\s+")


@dataclass
class Chapter:
    title: str
    paragraphs: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n\n".join(self.paragraphs)

    @property
    def word_count(self) -> int:
        return sum(len(p.split()) for p in self.paragraphs)


@dataclass
class Ingested:
    chapters: list[Chapter]
    page_count: int = 0

    @property
    def word_count(self) -> int:
        return sum(c.word_count for c in self.chapters)

    @property
    def full_text(self) -> str:
        return "\n\n".join(c.text for c in self.chapters)


def _clean(text: str) -> str:
    return _WS.sub(" ", text).strip()


def _paragraphs_from_html(html: str) -> list[str]:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "header", "footer"]):
        tag.decompose()
    paras: list[str] = []
    for el in soup.find_all(["p", "h1", "h2", "h3", "h4", "li", "blockquote"]):
        t = _clean(el.get_text(" ", strip=True))
        if t:
            paras.append(t)
    if not paras:
        body = _clean(soup.get_text(" ", strip=True))
        if body:
            paras = [body]
    return paras


def _chapter_title(soup_or_html: str | BeautifulSoup, fallback: str) -> str:
    soup = (
        soup_or_html
        if isinstance(soup_or_html, BeautifulSoup)
        else BeautifulSoup(soup_or_html, "html.parser")
    )
    for tag_name in ("h1", "h2", "title"):
        tag = soup.find(tag_name)
        if tag:
            t = _clean(tag.get_text(" ", strip=True))
            if t:
                return t
    return fallback


def ingest_epub(path: Path) -> Ingested:
    book = epub.read_epub(str(path), options={"ignore_ncx": True})
    chapters: list[Chapter] = []
    idx = 0
    for item in book.get_items():
        if item.get_type() != ebooklib.ITEM_DOCUMENT:
            continue
        html = item.get_content().decode("utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        paras = _paragraphs_from_html(html)
        if not paras:
            continue
        idx += 1
        title = _chapter_title(soup, f"Chapter {idx}")
        chapters.append(Chapter(title=title, paragraphs=paras))
    return Ingested(chapters=chapters, page_count=len(chapters))


def ingest_pdf(path: Path) -> Ingested:
    doc = pymupdf.open(str(path))
    paragraphs: list[str] = []
    for page in doc:
        text = page.get_text("text")
        for block in text.split("\n\n"):
            t = _clean(block)
            if t:
                paragraphs.append(t)
    page_count = doc.page_count
    doc.close()
    chapter = Chapter(title=path.stem, paragraphs=paragraphs)
    return Ingested(chapters=[chapter], page_count=page_count)


def ingest_txt(path: Path) -> Ingested:
    text = path.read_text(encoding="utf-8", errors="ignore")
    paragraphs = [_clean(b) for b in text.split("\n\n") if _clean(b)]
    chapter = Chapter(title=path.stem, paragraphs=paragraphs)
    return Ingested(chapters=[chapter], page_count=1)


def ingest_docx(path: Path) -> Ingested:
    doc = DocxDocument(str(path))
    chapters: list[Chapter] = []
    current: Chapter | None = None
    fallback_idx = 0
    for para in doc.paragraphs:
        text = _clean(para.text)
        if not text:
            continue
        style_name = (para.style.name if para.style else "") or ""
        style = style_name.lower()
        is_heading = style.startswith("heading") or style == "title"
        if is_heading:
            current = Chapter(title=text, paragraphs=[])
            chapters.append(current)
            continue
        if current is None:
            fallback_idx += 1
            current = Chapter(
                title=path.stem if fallback_idx == 1 else f"Section {fallback_idx}",
                paragraphs=[],
            )
            chapters.append(current)
        current.paragraphs.append(text)
    if not chapters:
        chapters = [Chapter(title=path.stem, paragraphs=[])]
    return Ingested(chapters=chapters, page_count=max(1, len(chapters)))


_EXT = {
    ".epub": ingest_epub,
    ".pdf": ingest_pdf,
    ".txt": ingest_txt,
    ".md": ingest_txt,
    ".docx": ingest_docx,
}


def ingest(path: Path) -> Ingested:
    ext = path.suffix.lower()
    fn = _EXT.get(ext)
    if fn is None:
        raise ValueError(f"unsupported file type: {ext}")
    return fn(path)


def supported_extensions() -> tuple[str, ...]:
    return tuple(_EXT.keys())
